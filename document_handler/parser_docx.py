import docx
import os
from typing import Dict, Any, Optional, List
from utils.logger import get_logger

logger = get_logger(__name__)

class DocxParser:
    """Parser for Microsoft Word DOCX documents"""
    
    def __init__(self):
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        
    def parse_docx(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Parse DOCX file and extract content and metadata"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"DOCX file not found: {file_path}")
                return None
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                logger.error(f"DOCX file too large: {file_size} bytes")
                return None
            
            doc = docx.Document(file_path)
            
            result = {
                'type': 'docx',
                'extracted_text': '',
                'metadata': {
                    'size': file_size,
                    'file_path': file_path,
                    'paragraphs': 0,
                    'tables': 0,
                    'images': 0
                },
                'paragraphs': [],
                'tables': [],
                'styles_used': [],
                'success': True
            }
            
            # Extract core properties
            core_props = doc.core_properties
            result['metadata'].update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or '',
                'version': core_props.version or ''
            })
            
            # Extract paragraphs
            full_text = ""
            styles_used = set()
            
            for para in doc.paragraphs:
                if para.text.strip():
                    para_data = {
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal',
                        'alignment': str(para.alignment) if para.alignment else 'Unknown'
                    }
                    
                    result['paragraphs'].append(para_data)
                    full_text += para.text + "\n"
                    
                    if para.style:
                        styles_used.add(para.style.name)
            
            result['extracted_text'] = full_text.strip()
            result['styles_used'] = list(styles_used)
            result['metadata']['paragraphs'] = len(result['paragraphs'])
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = self._extract_table_data(table, table_idx)
                result['tables'].append(table_data)
                
                # Add table content to extracted text
                full_text += f"\n--- Table {table_idx + 1} ---\n"
                for row in table_data['rows']:
                    full_text += " | ".join(row) + "\n"
            
            result['metadata']['tables'] = len(result['tables'])
            result['extracted_text'] = full_text.strip()
            
            # Count words and characters
            result['metadata']['total_characters'] = len(result['extracted_text'])
            result['metadata']['total_words'] = len(result['extracted_text'].split())
            
            # Analyze content
            content_analysis = self._analyze_docx_content(result['extracted_text'], result['styles_used'])
            result['content_analysis'] = content_analysis
            
            # Extract images info (basic)
            result['metadata']['images'] = self._count_images(doc)
            
            logger.info(f"Successfully parsed DOCX: {file_path}, {result['metadata']['paragraphs']} paragraphs, {result['metadata']['tables']} tables")
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return {
                'type': 'docx',
                'extracted_text': '',
                'metadata': {'error': str(e)},
                'success': False
            }
    
    def extract_tables(self, file_path: str) -> Optional[List[Dict[str, Any]]]:
        """Extract all tables from DOCX document"""
        try:
            result = self.parse_docx(file_path)
            if not result or not result['success']:
                return None
            
            return result['tables']
            
        except Exception as e:
            logger.error(f"Error extracting tables from DOCX: {e}")
            return None
    
    def extract_headings(self, file_path: str) -> Optional[List[Dict[str, str]]]:
        """Extract headings and their hierarchy"""
        try:
            doc = docx.Document(file_path)
            headings = []
            
            for para in doc.paragraphs:
                if para.style and 'Heading' in para.style.name:
                    headings.append({
                        'text': para.text,
                        'level': para.style.name,
                        'style': para.style.name
                    })
            
            return headings
            
        except Exception as e:
            logger.error(f"Error extracting headings from DOCX: {e}")
            return None
    
    def _extract_table_data(self, table, table_idx: int) -> Dict[str, Any]:
        """Extract data from a table"""
        table_data = {
            'table_index': table_idx,
            'rows': [],
            'columns': 0,
            'has_header': False
        }
        
        try:
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                
                table_data['rows'].append(row_data)
                
                # Update column count
                if len(row_data) > table_data['columns']:
                    table_data['columns'] = len(row_data)
            
            # Simple heuristic to detect header row
            if table_data['rows']:
                first_row = table_data['rows'][0]
                # If first row has mostly non-numeric content, it might be a header
                non_numeric_cells = sum(1 for cell in first_row if not self._is_mostly_numeric(cell))
                table_data['has_header'] = non_numeric_cells > len(first_row) / 2
                
        except Exception as e:
            logger.error(f"Error extracting table data: {e}")
        
        return table_data
    
    def _analyze_docx_content(self, text: str, styles_used: List[str]) -> Dict[str, Any]:
        """Analyze DOCX content to determine document type and characteristics"""
        import re
        
        analysis = {
            'document_type': 'unknown',
            'structure_score': 0.0,
            'content_type': 'unknown',
            'key_sections': [],
            'formal_level': 'unknown'
        }
        
        text_lower = text.lower()
        
        # Determine document type based on content and structure
        doc_type_indicators = {
            'contract': ['kontrak', 'perjanjian', 'pasal', 'ayat', 'syarat dan ketentuan', 'pihak pertama', 'pihak kedua'],
            'specification': ['spesifikasi', 'persyaratan', 'standar', 'mutu', 'kualitas', 'metode kerja'],
            'report': ['laporan', 'executive summary', 'kesimpulan', 'rekomendasi', 'analisis'],
            'proposal': ['proposal', 'penawaran', 'tender', 'lelang', 'bid'],
            'manual': ['panduan', 'petunjuk', 'prosedur', 'langkah-langkah', 'cara'],
            'correspondence': ['surat', 'memo', 'nota', 'undangan', 'pemberitahuan']
        }
        
        max_score = 0
        detected_type = 'unknown'
        
        for doc_type, indicators in doc_type_indicators.items():
            score = sum(1 for indicator in indicators if indicator in text_lower)
            if score > max_score:
                max_score = score
                detected_type = doc_type
        
        analysis['document_type'] = detected_type
        
        # Analyze structure based on styles used
        structure_indicators = ['Heading', 'Title', 'Subtitle', 'List', 'Caption', 'Quote']
        structure_score = sum(1 for style in styles_used if any(indicator in style for indicator in structure_indicators))
        analysis['structure_score'] = min(structure_score / 5.0, 1.0)  # Normalize to 0-1
        
        # Determine formality level
        formal_indicators = ['dengan hormat', 'yang bertanda tangan', 'demikian', 'terima kasih', 'hormat kami']
        formal_count = sum(1 for indicator in formal_indicators if indicator in text_lower)
        
        if formal_count >= 2:
            analysis['formal_level'] = 'formal'
        elif formal_count >= 1:
            analysis['formal_level'] = 'semi-formal'
        else:
            analysis['formal_level'] = 'informal'
        
        # Extract key sections based on common patterns
        section_patterns = [
            r'(?:^|\n)([A-Z][^.\n]{5,50})(?:\n|$)',  # Lines that might be section headers
            r'(?:^|\n)(\d+\.?\s+[A-Z][^.\n]{5,50})(?:\n|$)',  # Numbered sections
            r'(?:^|\n)([IVX]+\.?\s+[A-Z][^.\n]{5,50})(?:\n|$)'  # Roman numbered sections
        ]
        
        key_sections = []
        for pattern in section_patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            key_sections.extend(matches[:10])  # Limit to 10 sections
        
        analysis['key_sections'] = list(set(key_sections))[:10]
        
        return analysis
    
    def _count_images(self, doc) -> int:
        """Count images in the document"""
        try:
            image_count = 0
            
            # Count inline shapes (images)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            return image_count
            
        except Exception as e:
            logger.debug(f"Error counting images: {e}")
            return 0
    
    def _is_mostly_numeric(self, text: str) -> bool:
        """Check if text is mostly numeric"""
        if not text.strip():
            return False
        
        import re
        # Remove common non-numeric characters in numbers
        cleaned = re.sub(r'[.,\s%Rp-]', '', text)
        if not cleaned:
            return False
        
        numeric_chars = sum(1 for char in cleaned if char.isdigit())
        return numeric_chars / len(cleaned) > 0.5
